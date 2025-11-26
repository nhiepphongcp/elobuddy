from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from docxtpl import DocxTemplate # Correct import for python-docx-template
import os
import io
import datetime
import json

app = Flask(__name__)
CORS(app) # Keep CORS for API, though direct HTML serving might not need it as much for same-origin.

# --- Configuration for Data Files and Templates ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = BASE_DIR # Store JSON files in the same directory as app.py for simplicity
TEMPLATE_DOCX_DIR = os.path.join(BASE_DIR, 'templates_docx')
SERVED_HTML_DIR = os.path.join(BASE_DIR, 'served_html') # Directory to serve HTML files from

USERS_FILE = os.path.join(DATA_DIR, 'users.json')
DEPARTMENTS_FILE = os.path.join(DATA_DIR, 'departments.json')
REPAIR_REQUESTS_FILE = os.path.join(DATA_DIR, 'repair_requests.json')

PHIEU_YEU_CAU_TEMPLATE = os.path.join(TEMPLATE_DOCX_DIR, 'Phieu_Yeu_Cau_Sua_Chua.docx')
BIEN_BAN_KIEM_TRA_TEMPLATE = os.path.join(TEMPLATE_DOCX_DIR, 'Bien_Ban_Kiem_Tra (4).docx')

# --- Helper Functions for JSON Data Handling ---
def load_json_data(file_path, default_data=None):
    if default_data is None:
        default_data = [] if 'requests' in file_path or 'departments' in file_path else {}
    try:
        if not os.path.exists(file_path):
            save_json_data(file_path, default_data) # Create file with default if not exists
            return default_data
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        save_json_data(file_path, default_data) # Recreate with default if error
        return default_data

def save_json_data(file_path, data):
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    except IOError as e:
        print(f"Error saving data to {file_path}: {e}")


# --- Initialize Data Files and Directories ---
def initialize_default_data():
    if not os.path.exists(DATA_DIR): # Should already exist as BASE_DIR
        os.makedirs(DATA_DIR)

    if not os.path.exists(SERVED_HTML_DIR):
        os.makedirs(SERVED_HTML_DIR)
        print(f"Created directory for served HTML: {SERVED_HTML_DIR}")
        print(f"IMPORTANT: Please place 'admin_app.html' and 'department_app.html' into this '{SERVED_HTML_DIR}' directory.")


    if not os.path.exists(USERS_FILE):
        initial_users = {
            'admin': { 'id': 'admin', 'password': 'admin', 'role': 'admin', 'fullName': 'Admin Vật Tư', 'departmentName': 'Phòng Vật tư TBYT', 'departmentId': 'VT', 'email': 'admin.vt@benhvien.vn', 'phone': '0123456789' },
            'khoa_hstc': { 'id': 'khoa_hstc', 'password': '123', 'role': 'department', 'fullName': 'Nguyễn Văn An', 'departmentName': 'Khoa Hồi sức tích cực', 'departmentId': 'HSTC', 'email': 'van.an@benhvien.vn', 'phone': '0987654321' }
        }
        save_json_data(USERS_FILE, initial_users)

    if not os.path.exists(DEPARTMENTS_FILE):
        initial_departments = [
            { 'id': 'HSTC', 'name': 'Khoa Hồi sức tích cực', 'contactPerson': 'Bs. An', 'phone': '111', 'email': 'hstc@bv.vn' },
            { 'id': 'NOI', 'name': 'Khoa Nội tổng hợp', 'contactPerson': 'Bs. Bình', 'phone': '222', 'email': 'noi@bv.vn' },
            { 'id': 'VT', 'name': 'Phòng Vật tư TBYT', 'contactPerson': 'Admin VT', 'phone': '000', 'email': 'vt@bv.vn'}
        ]
        save_json_data(DEPARTMENTS_FILE, initial_departments)

    if not os.path.exists(REPAIR_REQUESTS_FILE):
        initial_repair_requests = []
        save_json_data(REPAIR_REQUESTS_FILE, initial_repair_requests)
    
    if not os.path.exists(TEMPLATE_DOCX_DIR):
        os.makedirs(TEMPLATE_DOCX_DIR)
        print(f"Created template directory: {TEMPLATE_DOCX_DIR}")
        if not os.path.exists(PHIEU_YEU_CAU_TEMPLATE):
            try:
                from docx import Document as DocxDocument # Renamed to avoid conflict
                doc = DocxDocument()
                doc.add_heading('PHIẾU YÊU CẦU SỬA CHỮA (Mẫu tạm)', 0)
                doc.add_paragraph('Khoa, Phòng: {{ khoa_phong }}')
                doc.save(PHIEU_YEU_CAU_TEMPLATE)
                print(f"Created dummy template: {PHIEU_YEU_CAU_TEMPLATE}")
            except ImportError:
                print("python-docx library not found, cannot create dummy DOCX template. Please install it: pip install python-docx")


        if not os.path.exists(BIEN_BAN_KIEM_TRA_TEMPLATE):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading('BIÊN BẢN KIỂM TRA (Mẫu tạm)', 0)
                doc.add_paragraph('Khoa, phòng: {{ bb_khoa_phong_header }}')
                doc.save(BIEN_BAN_KIEM_TRA_TEMPLATE)
                print(f"Created dummy template: {BIEN_BAN_KIEM_TRA_TEMPLATE}")
            except ImportError:
                 print("python-docx library not found, cannot create dummy DOCX template. Please install it: pip install python-docx")

initialize_default_data()

# --- Routes to serve HTML files ---
@app.route('/')
def serve_admin_app():
    if not os.path.exists(os.path.join(SERVED_HTML_DIR, 'admin_app.html')):
        return f"File 'admin_app.html' không tìm thấy trong thư mục '{SERVED_HTML_DIR}'. Vui lòng đặt file vào đúng vị trí.", 404
    return send_from_directory(SERVED_HTML_DIR, 'admin_app.html')

@app.route('/department')
def serve_department_app():
    if not os.path.exists(os.path.join(SERVED_HTML_DIR, 'department_app.html')):
        return f"File 'department_app.html' không tìm thấy trong thư mục '{SERVED_HTML_DIR}'. Vui lòng đặt file vào đúng vị trí.", 404
    return send_from_directory(SERVED_HTML_DIR, 'department_app.html')

# --- Route to serve static assets like images ---
@app.route('/assets/<path:filename>')
def serve_assets(filename):
    # This route will serve files from the 'served_html/assets' directory
    assets_dir = os.path.join(SERVED_HTML_DIR, 'assets')
    return send_from_directory(assets_dir, filename)


# --- API Endpoints for Data Management ---

# Users
@app.route('/api/users', methods=['GET'])
def get_users():
    users_data = load_json_data(USERS_FILE, {})
    return jsonify(users_data)

@app.route('/api/users', methods=['POST'])
def add_user():
    users_data = load_json_data(USERS_FILE, {})
    new_user = request.get_json()
    user_id = new_user.get('id')
    if not user_id:
        return jsonify({'error': 'User ID is required'}), 400
    if user_id in users_data:
        return jsonify({'error': 'User ID already exists'}), 400
    users_data[user_id] = new_user
    save_json_data(USERS_FILE, users_data)
    return jsonify({'message': 'User added successfully', 'user': new_user}), 201

@app.route('/api/users/<user_id>', methods=['POST']) 
def update_user(user_id):
    users_data = load_json_data(USERS_FILE, {})
    if user_id not in users_data:
        return jsonify({'error': 'User not found'}), 404
    updated_data = request.get_json()
    if 'password' not in updated_data or not updated_data['password']:
        updated_data['password'] = users_data[user_id].get('password')
    
    users_data[user_id].update(updated_data)
    save_json_data(USERS_FILE, users_data)
    return jsonify({'message': 'User updated successfully', 'user': users_data[user_id]})

@app.route('/api/users/<user_id>', methods=['DELETE'])
def delete_user(user_id):
    users_data = load_json_data(USERS_FILE, {})
    if user_id not in users_data:
        return jsonify({'error': 'User not found'}), 404
    
    repair_requests_data = load_json_data(REPAIR_REQUESTS_FILE, [])
    if any(req.get('requestingUserId') == user_id for req in repair_requests_data):
        return jsonify({'error': 'Không thể xóa người dùng vì còn yêu cầu sửa chữa liên kết.'}), 400

    del users_data[user_id]
    save_json_data(USERS_FILE, users_data)
    return jsonify({'message': 'User deleted successfully'})

# Departments
@app.route('/api/departments', methods=['GET'])
def get_departments():
    departments_data = load_json_data(DEPARTMENTS_FILE, [])
    return jsonify(departments_data)

@app.route('/api/departments', methods=['POST'])
def add_department():
    departments_data = load_json_data(DEPARTMENTS_FILE, [])
    new_dept = request.get_json()
    dept_id = new_dept.get('id')
    if not dept_id:
        return jsonify({'error': 'Department ID is required'}), 400
    if any(d.get('id') == dept_id for d in departments_data):
        return jsonify({'error': 'Department ID already exists'}), 400
    departments_data.append(new_dept)
    save_json_data(DEPARTMENTS_FILE, departments_data)
    return jsonify({'message': 'Department added successfully', 'department': new_dept}), 201

@app.route('/api/departments/<dept_id>', methods=['DELETE'])
def delete_department(dept_id):
    departments_data = load_json_data(DEPARTMENTS_FILE, [])
    users_data = load_json_data(USERS_FILE, {})
    repair_requests_data = load_json_data(REPAIR_REQUESTS_FILE, [])

    if dept_id == 'VT': 
        return jsonify({'error': 'Không thể xóa Phòng Vật tư TBYT.'}), 400
    
    if any(user.get('departmentId') == dept_id for user_id, user in users_data.items()):
        return jsonify({'error': 'Không thể xóa khoa/phòng vì còn tài khoản người dùng liên kết.'}), 400
    if any(req.get('departmentId') == dept_id for req in repair_requests_data):
        return jsonify({'error': 'Không thể xóa khoa/phòng vì còn yêu cầu sửa chữa liên kết.'}), 400

    departments_data = [d for d in departments_data if d.get('id') != dept_id]
    save_json_data(DEPARTMENTS_FILE, departments_data)
    return jsonify({'message': 'Department deleted successfully'})

# Repair Requests
@app.route('/api/repair_requests', methods=['GET'])
def get_repair_requests():
    repair_requests_data = load_json_data(REPAIR_REQUESTS_FILE, [])
    return jsonify(repair_requests_data)

@app.route('/api/repair_requests', methods=['POST'])
def add_repair_request():
    repair_requests_data = load_json_data(REPAIR_REQUESTS_FILE, [])
    new_request = request.get_json()
    
    new_id = 1
    if repair_requests_data:
        new_id = max(req.get('id', 0) for req in repair_requests_data) + 1
    new_request['id'] = new_id
    
    new_request.setdefault('receptionTimestamp', None)
    new_request.setdefault('completionTimestamp', None)
    new_request.setdefault('inspectionResult', '')
    new_request.setdefault('solutionDetails', '')
    new_request.setdefault('inspectionTimestamp', None)

    repair_requests_data.append(new_request)
    save_json_data(REPAIR_REQUESTS_FILE, repair_requests_data)
    return jsonify({'message': 'Repair request added successfully', 'request': new_request}), 201

@app.route('/api/repair_requests/<int:request_id>', methods=['POST']) 
def update_repair_request(request_id):
    repair_requests_data = load_json_data(REPAIR_REQUESTS_FILE, [])
    request_index = -1
    for i, req in enumerate(repair_requests_data):
        if req.get('id') == request_id:
            request_index = i
            break
    
    if request_index == -1:
        return jsonify({'error': 'Repair request not found'}), 404
        
    updated_data = request.get_json()
    repair_requests_data[request_index].update(updated_data)
    
    save_json_data(REPAIR_REQUESTS_FILE, repair_requests_data)
    return jsonify({'message': 'Repair request updated successfully', 'request': repair_requests_data[request_index]})


# --- DOCX Generation Endpoints ---
@app.route('/api/generate_phieu_yeu_cau', methods=['POST'])
def generate_phieu_yeu_cau():
    try:
        data = request.get_json()
        required_fields = ['khoa_phong', 'nguoi_thong_bao', 'ten_thiet_bi', 'ngay_gio_thong_bao', 'mo_ta_su_co']
        for field in required_fields:
            if field not in data or not data.get(field):
                if field not in ['truong_pho_khoa_phong', 'nguoi_nhan_thong_bao']:
                     return jsonify({'error': f'Thiếu trường thông tin bắt buộc: {field}'}), 400

        if not os.path.exists(PHIEU_YEU_CAU_TEMPLATE):
            return jsonify({'error': f'Không tìm thấy file mẫu: {os.path.basename(PHIEU_YEU_CAU_TEMPLATE)} tại {TEMPLATE_DOCX_DIR}'}), 500

        doc = DocxTemplate(PHIEU_YEU_CAU_TEMPLATE)
        context = {
            'khoa_phong': data.get('khoa_phong', ''),
            'truong_pho_khoa_phong': data.get('truong_pho_khoa_phong', '___________________________'),
            'nguoi_thong_bao': data.get('nguoi_thong_bao', ''),
            'ten_thiet_bi': data.get('ten_thiet_bi', ''),
            'ngay_gio_thong_bao': data.get('ngay_gio_thong_bao', ''),
            'nguoi_nhan_thong_bao': data.get('nguoi_nhan_thong_bao', '___________________________'),
            'mo_ta_su_co': data.get('mo_ta_su_co', '')
        }
        doc.render(context)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"PhieuYeuCauSuaChua_{data.get('ten_thiet_bi', 'ThietBi').replace(' ', '_')}_{timestamp}.docx"
        return send_file(
            file_stream, as_attachment=True, download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error in /api/generate_phieu_yeu_cau: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate_bien_ban_kiem_tra', methods=['POST'])
def generate_bien_ban_kiem_tra():
    try:
        data = request.get_json()
        required_fields = [
            'bb_khoa_phong_header', 'thoi_diem_yeu_cau', 'thoi_diem_kiem_tra', 
            'nguoi_quan_ly_su_dung', 'bb_khoa_phong_user', 'bb_ten_thiet_bi', 
            'ket_qua_kiem_tra', 'huong_khac_phuc_ket_qua', 
            'ten_khoa_su_dung_ky', 'ten_nguoi_kiem_tra_ky'
        ]
        for field in required_fields:
            if field not in data or not data.get(field):
                if field not in ['dien_thoai_lien_he', 'email_lien_he', 'thoi_gian_su_dung']:
                    return jsonify({'error': f'Thiếu trường thông tin bắt buộc: {field}'}), 400

        if not os.path.exists(BIEN_BAN_KIEM_TRA_TEMPLATE):
             return jsonify({'error': f'Không tìm thấy file mẫu: {os.path.basename(BIEN_BAN_KIEM_TRA_TEMPLATE)} tại {TEMPLATE_DOCX_DIR}'}), 500

        doc = DocxTemplate(BIEN_BAN_KIEM_TRA_TEMPLATE)
        context = {
            'bb_khoa_phong_header': data.get('bb_khoa_phong_header', ''),
            'thoi_diem_yeu_cau': data.get('thoi_diem_yeu_cau', ''),
            'thoi_diem_kiem_tra': data.get('thoi_diem_kiem_tra', ''),
            'nguoi_quan_ly_su_dung': data.get('nguoi_quan_ly_su_dung', ''),
            'bb_khoa_phong_user': data.get('bb_khoa_phong_user', ''),
            'dien_thoai_lien_he': data.get('dien_thoai_lien_he', '____________________'),
            'email_lien_he': data.get('email_lien_he', '____________________'),
            'bb_ten_thiet_bi': data.get('bb_ten_thiet_bi', ''),
            'thoi_gian_su_dung': data.get('thoi_gian_su_dung', '____________________'),
            'ket_qua_kiem_tra': data.get('ket_qua_kiem_tra', ''),
            'huong_khac_phuc_ket_qua': data.get('huong_khac_phuc_ket_qua', ''),
            'ten_khoa_su_dung_ky': data.get('ten_khoa_su_dung_ky', ''),
            'ten_nguoi_kiem_tra_ky': data.get('ten_nguoi_kiem_tra_ky', '')
        }
        doc.render(context)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"BienBanKiemTra_{data.get('bb_ten_thiet_bi', 'ThietBi').replace(' ', '_')}_{timestamp}.docx"
        return send_file(
            file_stream, as_attachment=True, download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error in /api/generate_bien_ban_kiem_tra: {e}")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True, ssl_context='adhoc')